import streamlit as st
import google.generativeai as genai
from youtube_transcript_api import YouTubeTranscriptApi
import yt_dlp
import os
import time
import re
import json
import tempfile
from docx import Document
from io import BytesIO
from google.api_core import exceptions

# --- 1. UIUX 專業樣式定義 ---
st.set_page_config(page_title="UA Sovereign Pro", page_icon="👑", layout="wide")

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;700;900&family=JetBrains+Mono&display=swap');
    html, body, [class*="css"] { font-family: 'Inter', sans-serif; background-color: #0D1117; color: #C9D1D9; }
    .main-title { font-size: 2.8rem; font-weight: 900; background: linear-gradient(90deg, #58a6ff, #bc8cff); -webkit-background-clip: text; -webkit-text-fill-color: transparent; }
    .report-card { background: #161b22; border: 1px solid #30363d; padding: 16px; border-radius: 10px; border-top: 3px solid #58a6ff; margin-bottom: 10px; height: 100%; }
    .shot-box { background: #161b22; border: 1px solid #30363d; border-radius: 12px; margin-bottom: 20px; overflow: hidden; }
    .ins-code { font-family: 'JetBrains Mono', monospace; font-size: 0.85rem; color: #7ee787; background: #010409; padding: 12px; border-radius: 8px; border: 1px solid #30363d; line-height: 1.4; word-break: break-all; }
</style>
""", unsafe_allow_html=True)

# --- 2. 核心工具函數 ---

def call_gemini_with_retry(model_name, contents, api_key, max_retries=3):
    if not api_key:
        st.error("❌ 找不到 API Key，請在側邊欄輸入。")
        return None
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(model_name)
    for i in range(max_retries):
        try:
            res = model.generate_content(contents)
            return res.text
        except exceptions.ResourceExhausted:
            time.sleep((i + 1) * 5)
        except Exception as e:
            st.error(f"❌ API 執行錯誤: {str(e)}")
            return None
    return None

def safe_json_parse(text):
    if not text: return {}
    # 移除 Markdown 的 JSON 標籤
    clean_text = re.sub(r'```json\s*|\s*```', '', text).strip()
    try:
        return json.loads(clean_text)
    except:
        # 嘗試擷取第一個 { 到最後一個 }
        try:
            match = re.search(r'\{.*\}', clean_text, re.DOTALL)
            if match: return json.loads(match.group())
        except: pass
    return {}

def create_docx_report(data, title="Report"):
    doc = Document()
    doc.add_heading(title, 0)
    if isinstance(data, dict):
        for k, v in data.items():
            doc.add_heading(str(k), level=1)
            doc.add_paragraph(str(v))
    else:
        doc.add_paragraph(str(data))
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def fetch_yt_info(url):
    try:
        # 更強大的 YouTube ID 提取
        regex = r"(?:v=|\/|shorts\/|be\/)([0-9A-Za-z_-]{11})"
        match = re.search(regex, url)
        if not match: return ""
        vid = match.group(1)
        
        with yt_dlp.YoutubeDL({'quiet': True, 'no_warnings': True}) as ydl:
            info = ydl.extract_info(url, download=False)
            title = info.get('title', 'Unknown Title')
            try:
                # 嘗試多種語言字幕
                transcript_list = YouTubeTranscriptApi.list_transcripts(vid)
                try:
                    t = transcript_list.find_transcript(['zh-TW', 'zh-HK', 'zh-Hans', 'en'])
                except:
                    t = transcript_list.find_generated_transcript(['zh-TW', 'zh', 'en'])
                
                script_data = t.fetch()
                script = " ".join([x['text'] for x in script_data])
            except: 
                script = "[無法獲取字幕，將僅參考標題與視覺內容]"
            return f"影片標題: {title}\n字幕內容: {script}"
    except Exception as e: 
        return f"解析 URL 失敗 ({url}): {str(e)}"

# --- 3. 初始化 Session ---
init_keys = {
    'step': 1, 'api_connected': False, 'api_key': "", 'active_model': "gemini-1.5-flash",
    'psy_analysis': {}, 'game_dna': {}, 'final_script': {}, 
    'ref_data_text': "", 'target_game': "Last War", 
    'duration': 30, 'style_preset': "死侍式賤萌", 'custom_style': "", 
    'yt_urls_raw': "", 'gemini_video_refs': []
}
for k, v in init_keys.items():
    if k not in st.session_state: st.session_state[k] = v

# --- 4. 側邊欄 ---
with st.sidebar:
    st.markdown("### 🛡️ 指揮中心")
    # 將 API Key 存入 session_state 確保全局可用
    st.session_state.api_key = st.text_input("輸入 Gemini API Key", value=st.session_state.api_key, type="password")
    
    if st.session_state.api_key:
        try:
            genai.configure(api_key=st.session_state.api_key)
            if not st.session_state.api_connected:
                # 獲取模型列表
                models = [m.name.split('/')[-1] for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
                st.session_state.active_model = st.selectbox("選擇 AI 模型", models, index=0)
                if st.button("🔗 測試並鎖定連線"):
                    test_res = call_gemini_with_retry(st.session_state.active_model, "Ping", st.session_state.api_key)
                    if test_res:
                        st.session_state.api_connected = True
                        st.success("✅ 連線成功")
                        st.rerun()
            else:
                st.success(f"● 已連線: {st.session_state.active_model}")
                if st.button("🔓 重新更換模型"):
                    st.session_state.api_connected = False
                    st.rerun()
        except Exception as e: 
            st.error(f"❌ Key 無效或網路錯誤: {e}")
    
    st.divider()
    if st.button("🔴 重置系統與清理雲端"):
        try:
            if st.session_state.api_connected:
                for f in genai.list_files(): genai.delete_file(f.name)
        except: pass
        for k in list(st.session_state.keys()): del st.session_state[k]
        st.rerun()

# --- 5. 主視覺導航 ---
st.markdown("<h1 class='main-title'>UA Sovereign Pro 👑</h1>", unsafe_allow_html=True)

nav_col = st.columns(4)
nav_labels = ["1. 素材注入", "2. 心理調研", "3. 戰術與 DNA", "4. 腳本生成"]
for i, label in enumerate(nav_labels):
    if nav_col[i].button(label, use_container_width=True, type="primary" if st.session_state.step == i+1 else "secondary"):
        st.session_state.step = i+1
        st.rerun()

TC_ONLY = "注意：除了生成的 Image/Motion Prompt 必須使用英文，其餘輸出務必使用『繁體中文 (zh-TW)』。嚴禁簡體。"

# --- STEP 1: 素材注入 ---
if st.session_state.step == 1:
    c1, c2 = st.columns(2)
    with c1:
        st.session_state.target_game = st.text_input("🎯 目標遊戲名稱", value=st.session_state.target_game)
    with c2:
        st.info("💡 提醒：分析參考片後，將提取其『成功邏輯』而非直接照抄物件。")

    st.session_state.yt_urls_raw = st.text_area("🔗 YouTube 連結 (一行一個)", value=st.session_state.yt_urls_raw, height=120)
    uploaded_videos = st.file_uploader("📁 上傳參考影片 (可多選)", accept_multiple_files=True, type=['mp4', 'mov', 'avi'])

    if st.button("🚀 第一步：啟動並行分析", type="primary", use_container_width=True):
        if not st.session_state.api_connected:
            st.error("請先在側邊欄連線 API"); st.stop()
        
        with st.status("正在注入多模態素材...", expanded=True) as status:
            # 修正變數名稱錯誤 (原本是 yt_urls)
            urls = [u.strip() for u in st.session_state.yt_urls_raw.split('\n') if u.strip()]
            ref_texts = []
            for u in urls:
                status.write(f"正在抓取 YouTube 字幕: {u}")
                ref_texts.append(fetch_yt_info(u))
            st.session_state.ref_data_text = "\n---\n".join(ref_texts)
            
            st.session_state.gemini_video_refs = []
            if uploaded_videos:
                for video in uploaded_videos:
                    status.write(f"正在上傳並處理影片: {video.name}")
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.mp4') as tmp:
                        tmp.write(video.read())
                        tmp_path = tmp.name
                    f_ref = genai.upload_file(tmp_path)
                    while f_ref.state.name == "PROCESSING":
                        time.sleep(2)
                        f_ref = genai.get_file(f_ref.name)
                    st.session_state.gemini_video_refs.append(f_ref)
                    os.remove(tmp_path)
            
            st.session_state.step = 2
            st.rerun()

# --- STEP 2: 心理調研 ---
elif st.session_state.step == 2:
    st.markdown("### 📊 素材心理基因報告")
    if not st.session_state.psy_analysis:
        with st.spinner("AI 正在觀看影片並解構戰略邏輯..."):
            contents = list(st.session_state.gemini_video_refs)
            prompt = f"""{TC_ONLY}
            請分析參考素材的『心理戰術框架』(禁止提到原片物件，只提取邏輯)。
            參考文字資訊：{st.session_state.ref_data_text}
            
            輸出格式請務必為純 JSON:
            {{
                "hook_logic": "前3秒如何吸引注意",
                "conflict_pattern": "衝突如何升級",
                "pacing_rhythm": "剪輯節奏描述",
                "visual_saliency": "視覺重心分佈",
                "reward_penalty": "正負回饋機制",
                "bartle_motivation": "玩家動機類型(成就/探索/社交/殺戮)",
                "sound_logic": "音效與音樂配合邏輯",
                "persuasion_arc": "如何說服玩家下載",
                "user_friction": "消除了哪些認知阻力",
                "emotional_hook": "喚起什麼情緒"
            }}"""
            psy_raw = call_gemini_with_retry(st.session_state.active_model, [prompt] + contents, st.session_state.api_key)
            st.session_state.psy_analysis = safe_json_parse(psy_raw)

    psy = st.session_state.psy_analysis
    if not psy:
        st.error("心理調研失敗，請重試。")
        if st.button("重新生成"): st.session_state.psy_analysis = {}; st.rerun()
    else:
        labels = {"hook_logic":"🪝 Hook 邏輯", "conflict_pattern":"⚔️ 衝突模式", "pacing_rhythm":"⏱️ 節奏律動", "visual_saliency":"👁️ 視覺權重", "reward_penalty":"🏆 獎懲邏輯", "bartle_motivation":"🧠 玩家動機", "sound_logic":"🎵 聲音設計", "persuasion_arc":"🎢 說服弧線", "user_friction":"🚧 用戶阻力", "emotional_hook":"🎭 情緒留存"}
        
        cols = st.columns(2)
        for i, (k, label) in enumerate(labels.items()):
            with cols[i % 2]:
                st.markdown(f"<div class='report-card'><h4>{label}</h4><p>{psy.get(k, 'N/A')}</p></div>", unsafe_allow_html=True)
        
        st.divider()
        docx_psy = create_docx_report({labels[k]: psy.get(k) for k in labels if k in psy}, title=f"{st.session_state.target_game} 心理戰研報告")
        st.download_button("📥 下載心理調研報告 (Word)", docx_psy, "Psy_Analysis.docx", use_container_width=True)
        
        if st.button("🚀 第二步：配置 DNA 與戰術", type="primary", use_container_width=True):
            st.session_state.step = 3; st.rerun()

# --- STEP 3: DNA 與風格校對 ---
elif st.session_state.step == 3:
    st.markdown("### 🧬 戰術配置中心：DNA、時長與風格")
    if not st.session_state.game_dna:
        with st.spinner("正在調研自家遊戲中..."):
            dna_prompt = f"{TC_ONLY}\n請分析遊戲《{st.session_state.target_game}》並輸出 JSON: genre(類型), description(介紹), objects(核心物件/角色清單), visual_style(視覺與畫質風格描述)"
            dna_raw = call_gemini_with_retry(st.session_state.active_model, dna_prompt, st.session_state.api_key)
            st.session_state.game_dna = safe_json_parse(dna_raw)

    dna = st.session_state.game_dna
    with st.form("dna_form_final"):
        c1, c2 = st.columns(2)
        dna['genre'] = c1.text_input("📍 遊戲類型", value=dna.get('genre', ''))
        st.session_state.duration = c2.slider("⏱️ 最終腳本秒數時長", 15, 60, st.session_state.duration, step=5)
        
        dna['description'] = st.text_area("📝 遊戲介紹與背景", value=dna.get('description', ''), height=80)
        dna['objects'] = st.text_area("📦 核心物件清單 (影響腳本出現的東西)", value=str(dna.get('objects', '')), height=100)
        dna['visual_style'] = st.text_area("🎨 視覺風格調性", value=dna.get('visual_style', ''), height=80)
        
        st.session_state.style_preset = st.selectbox("🎭 創意風格語氣", ["死侍式賤萌", "諾蘭式史詩", "極速流反轉", "日系吐槽", "自定義"])
        if st.session_state.style_preset == "自定義":
            st.session_state.custom_style = st.text_area("自定義風格詳細描述", value=st.session_state.custom_style)
            
        if st.form_submit_button("💾 儲存並準備生成", use_container_width=True):
            st.session_state.game_dna = dna; st.toast("戰略 DNA 已同步！")
    
    if st.button("🚀 生成最終原創腳本", type="primary", use_container_width=True):
        st.session_state.step = 4; st.rerun()

# --- STEP 4: 腳本生成 ---
elif st.session_state.step == 4:
    st.markdown("### 🎬 生成結果：跨維度原創腳本")
    final_tone = st.session_state.custom_style if st.session_state.style_preset == "自定義" else st.session_state.style_preset
    
    if st.button("🔄 重新生成腳本"): 
        st.session_state.final_script = {}
        st.rerun()

    if not st.session_state.final_script:
        with st.spinner(f"正在執行創意合成..."):
            prompt = f"""{TC_ONLY}
            請為《{st.session_state.target_game}》創作一則 UA 廣告腳本。
            畫風需求: {st.session_state.game_dna.get('visual_style')}。
            Image Prompt 必須包含 '--ar 9:16', 'high fidelity', 'unreal engine 5' 等關鍵字。
            
            時長: {st.session_state.duration} 秒
            風格語氣: {final_tone}
            參考心理框架: {json.dumps(st.session_state.psy_analysis, ensure_ascii=False)}
            遊戲核心物件: {json.dumps(st.session_state.game_dna.get('objects'), ensure_ascii=False)}
            
            輸出格式請務必為 JSON:
            {{ 
                "strategy_note": "戰略核心說明", 
                "script_steps": [ 
                    {{ 
                        "time_range": "00-03", 
                        "action": "視覺畫面描述", 
                        "dialogue": "旁白或台詞", 
                        "psychology_check": "此段對應的心理邏輯", 
                        "image_prompt": "給 Midjourney 的英文提示詞", 
                        "motion_prompt": "給 Luma/Runway 的英文動態提示詞" 
                    }} 
                ] 
            }}"""
            final_raw = call_gemini_with_retry(st.session_state.active_model, prompt, st.session_state.api_key)
            st.session_state.final_script = safe_json_parse(final_raw)

    fs = st.session_state.final_script
    if fs:
        st.success(f"💡 **戰略摘要**: {fs.get('strategy_note')}")
        
        script_text_for_word = f"遊戲：{st.session_state.target_game}\n創意戰略：{fs.get('strategy_note')}\n\n"
        
        for i, s in enumerate(fs.get('script_steps', [])):
            with st.container():
                st.markdown(f"""
                <div class='shot-box'>
                    <div style='background:#21262d; padding:12px 25px; display:flex; justify-content:space-between; align-items:center;'>
                        <b style="color: #58a6ff;">SHOT {i+1}</b> <span style='font-size:0.9rem; color:#8b949e;'>⏳ {s.get('time_range')}s</span>
                    </div>
                    <div style='padding:20px; display:grid; grid-template-columns: 1fr 1fr; gap:20px;'>
                        <div>
                            <div style='color:#ffa657; font-size:1.1rem; font-weight:bold; margin-bottom:8px;'>{s.get('action')}</div>
                            <div style='background:#0d1117; padding:15px; border-radius:8px; border-left:4px solid #bc8cff;'>{s.get('dialogue')}</div>
                            <p style='font-size:0.8rem; color:#8b949e; margin-top:10px;'>🎯 心理: {s.get('psychology_check')}</p>
                        </div>
                        <div>
                            <span style='font-size:0.75rem; font-weight:bold; color:#8b949e;'>🖼️ IMAGE PROMPT (Midjourney)</span>
                            <div class='ins-code'>{s.get('image_prompt')}</div>
                            <span style='font-size:0.75rem; font-weight:bold; color:#8b949e; margin-top:10px; display:block;'>🎥 MOTION PROMPT (Runway/Luma)</span>
                            <div class='ins-code' style='color:#79c0ff;'>{s.get('motion_prompt')}</div>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
            script_text_for_word += f"分鏡 {i+1} [{s.get('time_range')}s]: {s.get('action')}\n台詞：{s.get('dialogue')}\nImage Prompt: {s.get('image_prompt')}\nMotion Prompt: {s.get('motion_prompt')}\n\n"

        st.divider()
        docx_script = create_docx_report(script_text_for_word, title=f"{st.session_state.target_game} 分鏡腳本")
        st.download_button("📥 下載分鏡腳本 (Word)", docx_script, f"{st.session_state.target_game}_Script.docx", use_container_width=True)
    else:
        st.warning("未能生成有效 JSON 腳本，請點擊上方重新生成。")